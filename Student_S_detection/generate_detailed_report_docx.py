"""
Generate a detailed 5-6 page professional project report in DOCX format for the
Student Stress Detection System with comprehensive content.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os
from datetime import datetime


def set_run_font(run, font_name='Calibri', font_size=None, bold=False, italic=False, color=None):
    """Set font properties for a run"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph_with_formatting(doc, text, font_size=11, bold=False, italic=False, 
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=6):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold, italic=italic)
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = 1.15
    return para


def add_bullet_point(doc, text, font_size=11, spacing_after=4):
    """Add a bullet point"""
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_run_font(run, font_size=font_size)
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = 1.15
    return para


def add_image_to_doc(doc, image_path, width=5.5, caption=None):
    """Add an image to the document"""
    if os.path.exists(image_path):
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                set_run_font(caption_run, font_size=9, italic=True, color=(127, 140, 141))
                caption_para.paragraph_format.space_after = Pt(12)
            
            doc.add_paragraph()  # Add spacing after image
            return True
        except Exception as e:
            print(f"Error loading image {image_path}: {e}")
            return False
    return False


def create_performance_table(doc):
    """Create performance comparison table"""
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Model'
    header_cells[1].text = 'Accuracy (%)'
    
    # Set header formatting
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_size=11, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    data = [
        ['Logistic Regression', '89.11'],
        ['Support Vector Machine', '90.99'],
        ['Naive Bayes', '83.98'],
        ['Random Forest', '93.20'],
        ['Ensemble Model (Selected)', '91.36'],
    ]
    
    for i, (model, accuracy) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = model
        row_cells[1].text = accuracy
        
        # Highlight selected model (bold text)
        if 'Ensemble' in model:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, font_size=11, bold=True)
        
        # Center align accuracy column
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table


def build_docx(output_path="/home/sahu/Desktop/Student_S_detection/PROJECT_REPORT_DETAILED.docx"):
    """Build the detailed project report in DOCX format"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========== PAGE 1: TITLE AND ABSTRACT ==========
    # Title
    title = doc.add_heading('Student Stress Detection System', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, font_size=22, bold=True, color=(44, 62, 80))
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('Using Ensemble Machine Learning')
    set_run_font(subtitle_run, font_size=12, color=(127, 140, 141))
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Project Report | {datetime.now().strftime("%B %d, %Y")}')
    set_run_font(date_run, font_size=10, color=(127, 140, 141))
    
    doc.add_paragraph()  # Spacing
    
    # Abstract
    add_heading(doc, 'Abstract', level=1)
    abstract_text = (
        "This project presents a comprehensive machine learning system for detecting and "
        "predicting student stress levels using ensemble learning techniques. The system "
        "analyzes ten key features related to student lifestyle, academic performance, and "
        "personal circumstances to classify stress levels into three categories: low, medium, "
        "and high. Four individual machine learning models—Logistic Regression, Support Vector "
        "Machine (SVM), Naive Bayes, and Random Forest—were trained and evaluated. An ensemble "
        "model combining all four models using a Voting Classifier achieved 91.36% accuracy. "
        "The system includes a complete data pipeline from synthetic data generation to "
        "preprocessing, model training, evaluation, and deployment through a web application. "
        "The web interface provides real-time stress predictions along with personalized "
        "motivational suggestions to help students manage their stress levels effectively."
    )
    add_paragraph_with_formatting(doc, abstract_text, font_size=11, spacing_after=12)
    
    doc.add_page_break()
    
    # ========== PAGE 2: PROBLEM STATEMENT ==========
    add_heading(doc, '1. Problem Statement', level=1)
    
    problem_text = (
        "Student stress has become a critical concern in educational institutions worldwide. "
        "Academic pressure, financial constraints, social challenges, and workload management "
        "contribute significantly to student stress levels. Early detection and intervention "
        "can help prevent severe mental health issues and improve overall academic performance "
        "and well-being."
    )
    add_paragraph_with_formatting(doc, problem_text, font_size=11)
    
    add_heading(doc, '1.1 Challenges', level=2)
    challenges = [
        "Traditional stress assessment methods rely on self-reported questionnaires and clinical "
        "evaluations, which are time-consuming and subjective.",
        "Manual assessments may not capture real-time stress patterns or identify at-risk "
        "students early enough.",
        "Lack of automated, scalable solutions for stress detection in educational settings.",
        "Need for actionable insights and personalized recommendations based on individual "
        "student circumstances.",
    ]
    for challenge in challenges:
        add_bullet_point(doc, challenge, font_size=11)
    
    add_heading(doc, '1.2 Objectives', level=2)
    objectives = [
        "Develop a robust machine learning system capable of accurately predicting student "
        "stress levels based on lifestyle and academic features.",
        "Compare multiple classification algorithms to identify the most effective approach.",
        "Implement ensemble learning to improve prediction accuracy and robustness.",
        "Create a user-friendly web application for real-time stress prediction and intervention.",
        "Provide actionable insights through personalized recommendations.",
    ]
    for obj in objectives:
        add_bullet_point(doc, obj, font_size=11)
    
    doc.add_page_break()
    
    # ========== PAGE 3: METHODOLOGY ==========
    add_heading(doc, '2. Methodology', level=1)
    
    add_heading(doc, '2.1 Dataset', level=2)
    dataset_text = (
        "A synthetic dataset of 10,000 student records was generated using Python, incorporating "
        "realistic distributions and relationships between features. The dataset includes ten "
        "features: study hours, sleep hours, exercise hours, social activities, assignment "
        "deadlines, exam pressure, family support, financial stress, academic performance, and "
        "workload level. The target variable has three classes: low, medium, and high stress levels."
    )
    add_paragraph_with_formatting(doc, dataset_text, font_size=11)
    
    add_heading(doc, '2.2 Data Preprocessing', level=2)
    preprocessing_steps = [
        "Data Cleaning: Removal of missing values, duplicates, and outliers using Interquartile "
        "Range (IQR) method.",
        "Data Balancing: Application of SMOTE (Synthetic Minority Oversampling Technique) to "
        "address class imbalance.",
        "Feature Scaling: StandardScaler normalization to ensure features have zero mean and "
        "unit variance.",
        "Train-Test Split: 80-20 stratified split to maintain class distribution.",
    ]
    for step in preprocessing_steps:
        add_bullet_point(doc, step, font_size=11)
    
    add_heading(doc, '2.3 Model Selection', level=2)
    model_selection_text = (
        "Four machine learning algorithms were selected for comparison: Logistic Regression "
        "(interpretable baseline), Support Vector Machine (non-linear classification), Naive Bayes "
        "(fast probabilistic classifier), and Random Forest (captures complex interactions). "
        "An ensemble model using Voting Classifier with soft voting was created to combine "
        "predictions from all four models."
    )
    add_paragraph_with_formatting(doc, model_selection_text, font_size=11)
    
    doc.add_page_break()
    
    # ========== PAGE 4: MODEL JUSTIFICATION ==========
    add_heading(doc, '3. Model Selection and Justification', level=1)
    
    add_heading(doc, '3.1 Individual Models', level=2)
    
    models_info = [
        (
            "Logistic Regression",
            "Provides interpretable coefficients showing feature importance. Fast training and "
            "prediction with probabilistic outputs. Works well when features have linear "
            "relationships with the target variable."
        ),
        (
            "Support Vector Machine (SVM)",
            "Handles complex decision boundaries using kernel trick. Effective with "
            "high-dimensional data and provides good generalization. Sensitive to feature scaling, "
            "requiring StandardScaler preprocessing."
        ),
        (
            "Naive Bayes",
            "Extremely fast training and prediction. Based on Bayes' theorem with feature "
            "independence assumption. Provides probability estimates and works well even with "
            "limited data."
        ),
        (
            "Random Forest",
            "Captures complex non-linear relationships and feature interactions. Provides feature "
            "importance insights. Handles outliers well and reduces overfitting through ensemble "
            "of decision trees. Achieved highest individual model accuracy (93.20%)."
        ),
    ]
    
    for model_name, description in models_info:
        para = doc.add_paragraph()
        run = para.add_run(f"{model_name}: ")
        set_run_font(run, font_size=11, bold=True)
        run2 = para.add_run(description)
        set_run_font(run2, font_size=11)
        para.paragraph_format.space_after = Pt(6)
    
    add_heading(doc, '3.2 Ensemble Model (Selected)', level=2)
    ensemble_text = (
        "The Voting Classifier ensemble model was selected as the best model for deployment. "
        "It combines all four individual models using soft voting, which averages probability "
        "predictions from each model. This approach offers several advantages: (1) Improved "
        "generalization by combining diverse learning approaches, (2) Reduced overfitting risk "
        "compared to individual models, (3) Better robustness to data variations, and (4) "
        "Probability estimates for uncertainty quantification. While Random Forest achieved "
        "slightly higher accuracy (93.20%), the ensemble model provides better balance between "
        "accuracy and generalization, making it more suitable for real-world deployment."
    )
    add_paragraph_with_formatting(doc, ensemble_text, font_size=11)
    
    doc.add_page_break()
    
    # ========== PAGE 5: RESULTS ==========
    add_heading(doc, '4. Results and Analysis', level=1)
    
    add_heading(doc, '4.1 Model Performance', level=2)
    performance_text = (
        "All models were evaluated on a held-out test set. The following table summarizes "
        "the accuracy achieved by each model:"
    )
    add_paragraph_with_formatting(doc, performance_text, font_size=11)
    
    # Add performance table
    create_performance_table(doc)
    doc.add_paragraph()  # Spacing after table
    
    add_heading(doc, '4.2 Performance Analysis', level=2)
    analysis_text = (
        "The ensemble model achieved 91.36% accuracy, demonstrating strong performance for a "
        "three-class classification problem. Analysis of the confusion matrix reveals excellent "
        "separation between low and medium stress levels, with some confusion between medium "
        "and high stress categories. The model provides probability distributions for each "
        "stress level, enabling uncertainty quantification and better decision-making."
    )
    add_paragraph_with_formatting(doc, analysis_text, font_size=11)
    
    # Add accuracy comparison image
    img_path = "/home/sahu/Desktop/Student_S_detection/evaluation/accuracy_comparison.png"
    add_image_to_doc(doc, img_path, width=5.5, caption="Figure 1: Model Accuracy Comparison")
    
    # Add confusion matrices image
    img_path2 = "/home/sahu/Desktop/Student_S_detection/evaluation/confusion_matrices.png"
    add_image_to_doc(doc, img_path2, width=5.5, caption="Figure 2: Confusion Matrices for All Models")
    
    doc.add_page_break()
    
    # ========== PAGE 6: FEATURE ANALYSIS AND CONCLUSION ==========
    add_heading(doc, '4.3 Feature Importance', level=2)
    feature_text = (
        "Based on the stress score calculation and model analysis, the most important features "
        "for stress prediction are: sleep hours (highest impact, inverse correlation), exam "
        "pressure, financial stress, assignment deadlines, and workload level. Protective "
        "factors include exercise hours and social activities (negative correlation with stress), "
        "while family support also reduces stress levels."
    )
    add_paragraph_with_formatting(doc, feature_text, font_size=11)
    
    # Add feature distribution or correlation image if available
    img_path3 = "/home/sahu/Desktop/Student_S_detection/evaluation/correlation_heatmap.png"
    add_image_to_doc(doc, img_path3, width=5.5, caption="Figure 3: Feature Correlation Heatmap")
    
    add_heading(doc, '5. Conclusion', level=1)
    
    conclusion_text = (
        "This project successfully developed a comprehensive student stress detection system "
        "using ensemble machine learning. The ensemble model achieved 91.36% accuracy, "
        "demonstrating strong performance for stress level classification. The system provides "
        "a complete end-to-end solution from data generation to web deployment, with "
        "personalized recommendations for stress management."
    )
    add_paragraph_with_formatting(doc, conclusion_text, font_size=11)
    
    add_heading(doc, '5.1 Key Achievements', level=2)
    achievements = [
        "Developed robust ensemble model achieving 91.36% accuracy",
        "Comprehensive comparison of four different ML algorithms",
        "Complete pipeline from data generation to web deployment",
        "User-friendly web application with real-time predictions",
        "Actionable insights through personalized recommendations",
    ]
    for achievement in achievements:
        add_bullet_point(doc, achievement, font_size=11)
    
    add_heading(doc, '5.2 Future Work', level=2)
    future_work = [
        "Validate on real-world student data from educational institutions",
        "Add temporal analysis to track stress levels over time",
        "Integrate with Learning Management Systems for automated data collection",
        "Develop mobile application for easier access",
        "Implement deep learning models for complex pattern recognition",
    ]
    for item in future_work:
        add_bullet_point(doc, item, font_size=11)
    
    # Save document
    doc.save(output_path)
    print(f"Detailed project report (DOCX) generated: {output_path}")
    print(f"Document contains {len(doc.paragraphs)} paragraphs")


if __name__ == "__main__":
    build_docx()

